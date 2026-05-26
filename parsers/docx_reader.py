"""
DOCX text extraction module for Zecpath AI hiring platform.
Uses python-docx to extract text from paragraphs, tables, headers, and footers.
"""

import os
from docx import Document
from docx.oxml.ns import qn
from utils.logger import get_logger

logger = get_logger(__name__)


class DOCXReader:
    """
    Extracts raw text from DOCX resumes using python-docx.
    Covers paragraphs, tables, headers, and footers.
    """

    def extract_text(self, file_path: str) -> str:
        """
        Extract raw text from a DOCX file.

        Args:
            file_path: Absolute or relative path to the DOCX file.

        Returns:
            Raw extracted text as a single string with sections joined by newlines.

        Raises:
            FileNotFoundError: If the file does not exist at the given path.
            ValueError: If the file is corrupted or cannot be read.
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(
                f"DOCX file not found at path: {file_path}"
            )

        all_text_blocks = []

        try:
            doc = Document(file_path)

            # Extract from paragraphs
            paragraph_text = self._extract_paragraphs(doc)
            if paragraph_text:
                all_text_blocks.append(paragraph_text)

            # Extract from tables
            table_text = self._extract_tables(doc)
            if table_text:
                all_text_blocks.append(table_text)

            # Extract from headers and footers
            header_footer_text = self._extract_headers_footers(doc)
            if header_footer_text:
                all_text_blocks.append(header_footer_text)

            logger.info(
                f"DOCX extraction successful: {os.path.basename(file_path)}"
            )

        except FileNotFoundError:
            raise
        except Exception as e:
            raise ValueError(
                f"Failed to read DOCX file '{file_path}'. "
                f"File may be corrupted or in an unsupported format. Error: {e}"
            )

        return "\n".join(all_text_blocks)

    def _extract_paragraphs(self, doc: Document) -> str:
        """
        Extract text from all paragraphs in the document body.

        Args:
            doc: An open python-docx Document object.

        Returns:
            Paragraph text joined by newlines.
        """
        lines = [para.text for para in doc.paragraphs if para.text.strip()]
        return "\n".join(lines)

    def _extract_tables(self, doc: Document) -> str:
        """
        Extract text from all tables in the document.

        Args:
            doc: An open python-docx Document object.

        Returns:
            Table cell text joined by newlines.
        """
        table_lines = []
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    table_lines.append("  ".join(row_cells))
        return "\n".join(table_lines)

    def _extract_headers_footers(self, doc: Document) -> str:
        """
        Extract text from headers and footers of all sections.

        Args:
            doc: An open python-docx Document object.

        Returns:
            Header and footer text joined by newlines.
        """
        hf_lines = []
        for section in doc.sections:
            for header in [section.header, section.first_page_header]:
                if header:
                    for para in header.paragraphs:
                        if para.text.strip():
                            hf_lines.append(para.text.strip())

            for footer in [section.footer, section.first_page_footer]:
                if footer:
                    for para in footer.paragraphs:
                        if para.text.strip():
                            hf_lines.append(para.text.strip())

        return "\n".join(hf_lines)